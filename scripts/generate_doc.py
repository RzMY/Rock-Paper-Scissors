"""
Generate a comprehensive project documentation Word document.
Run: python scripts/generate_doc.py
Output: docs/Rock-Paper-Scissors-项目文档.docx
"""
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime

BASE_DIR = Path(__file__).resolve().parent.parent
OUTPUT_DIR = BASE_DIR / "docs"
OUTPUT_DIR.mkdir(exist_ok=True)
OUTPUT_PATH = OUTPUT_DIR / "Rock-Paper-Scissors-项目文档.docx"

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hfont = hs.font
    hfont.name = '微软雅黑'
    hfont.color.rgb = RGBColor(0x0A, 0x0A, 0x14)
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def add_code_block(doc, code_text):
    """Add a code block styled paragraph."""
    for line in code_text.strip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(16)
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    return table

# ════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Rock-Paper-Scissors YOLO')
run.font.size = Pt(32)
run.font.bold = True
run.font.color.rgb = RGBColor(0xE9, 0x45, 0x60)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('基于 YOLOv8 手势识别的实时猜拳对战系统')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x88, 0x88, 0xAA)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run(f'文档版本: 1.0\n生成日期: {datetime.date.today().isoformat()}\n'
                    '技术栈: Python 3.10+ / FastAPI / YOLOv8 / WebSocket / Vanilla JS')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x5A, 0x5A, 0x7A)

doc.add_page_break()

# ════════════════════════════════════════════
# TABLE OF CONTENTS placeholder
# ════════════════════════════════════════════
doc.add_heading('目录', level=1)
toc_items = [
    '1. 项目概述', '2. 系统架构', '3. 目录结构',
    '4. 数据准备', '5. 模型训练', '6. 后端服务',
    '7. AI 策略引擎', '8. 游戏状态机', '9. WebSocket 协议',
    '10. REST API 参考', '11. 前端界面', '12. 数据库设计',
    '13. 配置文件', '14. 部署与运行', '15. 依赖清单',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ════════════════════════════════════════════
# 1. 项目概述
# ════════════════════════════════════════════
doc.add_heading('1. 项目概述', level=1)

doc.add_paragraph(
    'Rock-Paper-Scissors YOLO 是一个基于计算机视觉的实时猜拳（剪刀石头布）对战系统。'
    '系统使用 YOLOv8n (Nano) 目标检测模型对摄像头画面中的手势进行实时识别，'
    '玩家通过摄像头做出"石头"、"剪刀"或"布"手势，系统检测后与 AI 对手进行猜拳对决。'
)

doc.add_paragraph(
    'AI 对手采用四策略集成学习方案：马尔可夫链模式匹配、指数衰减频率分析、'
    '旋转周期检测以及元策略选择器，能够根据玩家的出拳历史自适应选择最优预测策略。'
)

doc.add_heading('核心特性', level=2)
features = [
    '实时手势识别：基于 YOLOv8n 模型，使用自定义数据集训练，支持三种手势分类',
    'AI 集成策略对手：4 种预测策略 + ε-greedy 元学习器，动态适应玩家习惯',
    '双向 WebSocket 通信：摄像头帧上传、标注回传、游戏状态同步均通过单 WebSocket',
    '暗色主题 Web UI：三栏布局（摄像头 / 游戏 / 统计），纯原生 JavaScript，无框架依赖',
    '持久化对局记录：SQLite 保存每轮结果，支持胜率/策略准确率统计',
    '自动对战模式：一键切换手动/自动，自动模式下连续对局',
]
for f in features:
    doc.add_paragraph(f, style='List Bullet')

doc.add_page_break()

# ════════════════════════════════════════════
# 2. 系统架构
# ════════════════════════════════════════════
doc.add_heading('2. 系统架构', level=1)

doc.add_paragraph(
    '系统采用经典的前后端分离架构，通过 WebSocket 维持低延迟的双向实时通信：'
)

doc.add_heading('架构图', level=2)
doc.add_paragraph(
    '┌─────────────────────────────────────────────────────────────────┐\n'
    '│                         Browser (Frontend)                       │\n'
    '│  ┌──────────┐  ┌──────────┐  ┌──────────┐  ┌──────────────┐   │\n'
    '│  │ Webcam   │  │ Game UI  │  │ AI       │  │ Stats Panel  │   │\n'
    '│  │ Manager  │  │ (State)  │  │ Console  │  │ (Gauge/Bars) │   │\n'
    '│  └────┬─────┘  └────┬─────┘  └──────────┘  └──────┬───────┘   │\n'
    '│       │             │                              │            │\n'
    '│       └──────┬──────┴──────────────────────────────┘            │\n'
    '│              │  WebSocket (ws://host:8000/ws)                    │\n'
    '└──────────────┼──────────────────────────────────────────────────┘\n'
    '               │\n'
    '┌──────────────┼──────────────────────────────────────────────────┐\n'
    '│              │           FastAPI Server                          │\n'
    '│  ┌───────────┴──────────┐  ┌─────────────┐  ┌──────────────┐   │\n'
    '│  │   WebSocket Handler  │  │  REST API   │  │  Static File │   │\n'
    '│  │   - frame processing │  │  /api/*     │  │  / (SPA)     │   │\n'
    '│  │   - game state mgmt  │  │             │  │              │   │\n'
    '│  └───────────┬──────────┘  └─────────────┘  └──────────────┘   │\n'
    '│              │                                                   │\n'
    '│  ┌───────────┼──────────────────────────────────────────────┐   │\n'
    '│  │           ▼                                               │   │\n'
    '│  │  ┌──────────┐  ┌──────────────┐  ┌────────────────────┐  │   │\n'
    '│  │  │ YOLO     │  │ Game Engine  │  │ AI Strategy        │  │   │\n'
    '│  │  │ Infer    │  │ (State Mach) │  │ (4-Strategy Ensem) │  │   │\n'
    '│  │  └──────────┘  └──────┬───────┘  └────────┬───────────┘  │   │\n'
    '│  │                       │                    │              │   │\n'
    '│  │                       ▼                    ▼              │   │\n'
    '│  │               ┌──────────────────────────────────┐        │   │\n'
    '│  │               │        SQLite Database            │        │   │\n'
    '│  │               │        (game_history.db)          │        │   │\n'
    '│  │               └──────────────────────────────────┘        │   │\n'
    '│  └───────────────────────────────────────────────────────────┘   │\n'
    '└──────────────────────────────────────────────────────────────────┘'
)

doc.add_heading('数据流', level=2)
doc.add_paragraph(
    '1. 浏览器通过 getUserMedia 捕获摄像头画面 → 每 67ms (~15 FPS) 截取一帧\n'
    '2. 前端将帧以 base64 JPEG 编码通过 WebSocket 发送 (type: "frame")\n'
    '3. 服务端 YOLO 推理 → 返回标注帧 + 检测结果列表\n'
    '4. 游戏状态机通过 tick() 驱动倒计时, 每秒推送 10 次状态更新\n'
    '5. 当状态为 SHOOT 时, 下一帧的最佳检测结果用于判定本回合胜负\n'
    '6. AI 策略引擎预测玩家出拳 → 计算机打出克制拳 → 结果写入 SQLite\n'
    '7. 回合结果广播至前端 → 更新比分、AI 控制台、统计面板'
)

doc.add_page_break()

# ════════════════════════════════════════════
# 3. 目录结构
# ════════════════════════════════════════════
doc.add_heading('3. 目录结构', level=1)

add_code_block(doc, """
Rock-Paper-Scissors/
├── backend/                    # 后端 Python 包
│   ├── __init__.py
│   ├── app.py                  # FastAPI 入口 + WebSocket + REST API
│   ├── config.py               # 全局配置常量
│   ├── model.py                # YOLO 推理封装 + 帧标注
│   ├── game_engine.py          # 游戏状态机 + 回合裁决
│   ├── ai_strategy.py          # AI 四策略集成引擎
│   ├── database.py             # SQLite 数据库操作层
│   └── static/                 # 前端静态资源
│       ├── index.html          # 主页面 (SPA)
│       ├── css/style.css       # 暗色主题样式
│       └── js/
│           ├── websocket.js    # WebSocket 客户端封装
│           ├── main.js         # App 入口 + 摄像头管理
│           ├── game.js         # 游戏 UI + 状态渲染
│           └── stats.js        # 统计面板 + 胜率仪表盘
├── scripts/                    # 工具脚本
│   ├── prepare_data.py         # 数据预处理 (Roboflow CSV → YOLO txt)
│   ├── train.py                # YOLOv8 训练脚本
│   └── generate_doc.py         # 本文档生成脚本
├── datasets/                   # 标注数据 (train/val/test)
│   ├── train/images/  + labels/
│   ├── val/images/    + labels/
│   └── test/images/   + labels/
├── models/                     # 训练产出模型文件
│   └── rps_yolov8n.pt          # 训练后的权重 (由 train.py 生成)
├── runs/                       # YOLO 训练运行输出
├── docs/                       # 文档输出目录
├── data.yaml                   # YOLO 数据集配置
├── requirements.txt            # Python 依赖
├── CLAUDE.md                   # Claude Code 项目指引
└── .gitignore
""")

doc.add_page_break()

# ════════════════════════════════════════════
# 4. 数据准备
# ════════════════════════════════════════════
doc.add_heading('4. 数据准备', level=1)

doc.add_paragraph(
    '数据集来源为 Roboflow 平台标注的手势数据。原始数据包含 Pascal VOC 格式的 '
    'CSV 标注文件 (_annotations.csv)，需要转换为 YOLO 训练所需的 TXT 格式。'
)

doc.add_heading('标注格式转换', level=2)
doc.add_paragraph(
    '输入 (Roboflow CSV): 每行包含 filename, class, xmin, ymin, xmax, ymax\n'
    '输出 (YOLO TXT):   每行 class_id cx cy w h (归一化到 0-1, 基于 640×640 图像)'
)

doc.add_heading('数据集划分', level=2)
doc.add_paragraph(
    '使用 sklearn 的 train_test_split (stratify 分层抽样, random_state=42):\n'
    '• 训练集: 原 train 目录的 80%\n'
    '• 验证集: 原 train 目录的 20%\n'
    '• 测试集: 直接使用原 test 目录'
)

doc.add_heading('类别定义', level=2)
add_table(doc,
    ['Class ID', '类别名称', '说明'],
    [
        ['0', 'Rock (石头)', '握拳'],
        ['1', 'Paper (布)', '手掌张开'],
        ['2', 'Scissors (剪刀)', '食指中指伸出'],
    ]
)

doc.add_heading('运行命令', level=2)
add_code_block(doc, '# 预处理数据集 (仅需执行一次)\npython scripts/prepare_data.py')

doc.add_page_break()

# ════════════════════════════════════════════
# 5. 模型训练
# ════════════════════════════════════════════
doc.add_heading('5. 模型训练', level=1)

doc.add_heading('模型选型', level=2)
doc.add_paragraph(
    '选用 Ultralytics YOLOv8n (Nano)，参数量约 3.2M，适合实时推理。\n'
    '输入分辨率: 640×640，与摄像头捕获分辨率一致。'
)

doc.add_heading('训练参数', level=2)
add_table(doc,
    ['参数', '默认值', '说明'],
    [
        ['--epochs', '1 (CI 验证) / 100 (正式训练)', '训练轮数'],
        ['--model', 'yolov8n.pt', '基础预训练模型'],
        ['--batch', '8', '批次大小'],
        ['--device', 'cuda', '训练设备 (cuda/cpu)'],
        ['--workers', '0', 'DataLoader 进程数 (Windows 安全默认)'],
        ['imgsz', '640', '输入图像尺寸'],
        ['amp', 'False', '自动混合精度 (Windows 环境关闭)'],
    ]
)

doc.add_heading('数据增强策略', level=2)
add_table(doc,
    ['增强方法', '参数', '作用'],
    [
        ['HSV 色相', 'hsv_h=0.015', '微调色调变化'],
        ['HSV 饱和度', 'hsv_s=0.7', '大幅调节饱和度'],
        ['HSV 明度', 'hsv_v=0.4', '调节亮度'],
        ['随机旋转', 'degrees=15°', '模拟不同手部角度'],
        ['平移', 'translate=0.1', '模拟手部位置偏移'],
        ['缩放', 'scale=0.5', '模拟远近变化'],
        ['水平翻转', 'fliplr=0.5', '50% 概率左右镜像'],
        ['Mosaic', 'mosaic=1.0', '四图拼接训练'],
        ['MixUp', 'mixup=0.1', '10% 概率图像混合'],
    ]
)

doc.add_heading('运行命令', level=2)
add_code_block(doc, """# GPU 训练 (100 epoch)
python scripts/train.py --epochs 100 --batch 64 --device cuda

# CPU 训练 (小批次快速验证)
python scripts/train.py --epochs 10 --batch 4 --device cpu

# 训练完成后最佳模型自动复制到 models/rps_yolov8n.pt""")

doc.add_page_break()

# ════════════════════════════════════════════
# 6. 后端服务
# ════════════════════════════════════════════
doc.add_heading('6. 后端服务', level=1)

doc.add_paragraph(
    '后端基于 FastAPI 构建，包含 REST API 和 WebSocket 两个通信通道。'
    '启动时自动加载 YOLO 模型（如果存在）并初始化游戏引擎和 SQLite 数据库。'
)

doc.add_heading('启动流程', level=2)
doc.add_paragraph(
    '1. FastAPI startup 事件触发\n'
    '2. 检查 models/rps_yolov8n.pt 是否存在 → 加载 YOLO 模型\n'
    '   • 存在: 正常加载，启用手势检测功能\n'
    '   • 不存在: 打印警告，服务正常运行但检测功能禁用\n'
    '3. 初始化 GameEngine（AI 策略、状态机）\n'
    '4. 挂载静态文件目录 → SPA 可通过 "/" 直接访问\n'
    '5. 监听 0.0.0.0:8000'
)

doc.add_heading('核心模块依赖关系', level=2)
doc.add_paragraph(
    'app.py\n'
    '  ├── config.py        → MODEL_PATH, WS_FRAME_QUALITY\n'
    '  ├── model.py         → YOLOInference.predict(), annotate_frame()\n'
    '  ├── game_engine.py   → GameEngine, GameState 状态机\n'
    '  └── database.py      → insert_round(), get_stats(), get_recent_rounds()'
)

doc.add_heading('运行命令', level=2)
add_code_block(doc, """# 方式一：直接运行
python backend/app.py

# 方式二：uvicorn 命令行
uvicorn backend.app:app --host 0.0.0.0 --port 8000 --reload""")

doc.add_page_break()

# ════════════════════════════════════════════
# 7. AI 策略引擎
# ════════════════════════════════════════════
doc.add_heading('7. AI 策略引擎', level=1)

doc.add_paragraph(
    'AI 策略引擎 (backend/ai_strategy.py) 采用四策略集成架构，'
    '通过元学习器动态选择当前最优的预测策略。所有策略共同评估，'
    '但每回合由元策略选择其中一个执行。'
)

doc.add_heading('策略一：马尔可夫链预测器 (MarkovChainPredictor)', level=2)
doc.add_paragraph(
    '• 原理: 基于 N-gram (默认 order=3) 模式匹配。统计历史中每种出拳序列的后继出拳分布\n'
    '• 训练: 每回合前用完整历史重新训练，滑动窗口扫描所有 N-1 元组\n'
    '• 预测: 取最近 N-1 次出拳 → 查表 → 返回最高频的后继出拳\n'
    '• fallback: 历史不足或找不到匹配模式时返回 None（不出候选）'
)

doc.add_heading('策略二：指数衰减频率分析器 (FrequencyAnalyzer)', level=2)
doc.add_paragraph(
    '• 原理: 对历史出拳进行加权计数，越近的出拳权重越高\n'
    '• 权重公式: weight = decay^(n-1-i)，其中 decay = 0.85\n'
    '• 特点: 始终返回预测结果（无 fallback），作为可靠后备策略\n'
    '• 示例: 最近 10 次出拳中，最近一次权重 = 0.85^0 = 1.0，10 次前权重 = 0.85^9 ≈ 0.23'
)

doc.add_heading('策略三：旋转周期检测器 (AntiRotator)', level=2)
doc.add_paragraph(
    '• 原理: 检测玩家是否有循环出拳习惯 (石头→布→剪刀→石头...)\n'
    '• 窗口: 最近 6 次出拳\n'
    '• 判定: 连续转移中顺时针或逆时针占比 ≥ 70% 时触发\n'
    '• 预测: 按检测到的旋转方向推断下一次出拳\n'
    '• fallback: 无明确旋转规律时返回 None'
)

doc.add_heading('策略零：元策略选择器 (MetaStrategy)', level=2)
doc.add_paragraph(
    '• 数据: 为每个子策略维护长度为 20 的滑动窗口正确率队列\n'
    '• 选择算法 (ε-greedy):\n'
    '  - 前 30 回合: 20% 概率随机探索，80% 概率选择准确率最高的策略\n'
    '  - 30 回合后: 总是选取准确率最高的策略\n'
    '• 评分: 每回合结束后，对所有产出候选预测的策略都进行评分（不限于被选中的策略）\n'
    '• 未测试策略: 默认 0.5 乐观先验'
)

doc.add_heading('AIStrategy 主控流程', level=2)
doc.add_paragraph(
    '1. predict_player_move(history):\n'
    '   a. 用完整历史训练 Markov 链\n'
    '   b. 调用三个子策略获取各自的预测\n'
    '   c. 元策略从可用候选中选择一个\n'
    '   d. 返回 (预测出拳, 策略名, 推理详情)\n\n'
    '2. get_computer_move(history):\n'
    '   调用 predict_player_move → 返回能击败预测出拳的招数\n\n'
    '3. record_result(strategy, predicted, actual, all_candidates):\n'
    '   记录被选中策略的表现，同时也记录其他策略的"反事实"表现'
)

doc.add_heading('胜负规则', level=2)
add_table(doc,
    ['玩家出拳', '电脑出拳', '结果', '说明'],
    [
        ['Rock', 'Paper', 'lose', '布包石头'],
        ['Rock', 'Scissors', 'win', '石头砸剪刀'],
        ['Paper', 'Scissors', 'lose', '剪刀剪纸'],
        ['Paper', 'Rock', 'win', '布包石头'],
        ['Scissors', 'Rock', 'lose', '石头砸剪刀'],
        ['Scissors', 'Paper', 'win', '剪刀剪纸'],
        ['任意', '相同', 'draw', '平局'],
    ]
)

doc.add_heading('关键配置', level=2)
add_table(doc,
    ['常量', '值', '说明'],
    [
        ['MARKOV_ORDER', '3', '马尔可夫链 N-gram 阶数'],
        ['FREQUENCY_DECAY', '0.85', '频率分析的指数衰减因子'],
        ['RECENT_WINDOW', '10', '推理中显示的历史长度'],
        ['META_WINDOW', '20', '元策略正确率滑动窗口大小'],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════
# 8. 游戏状态机
# ════════════════════════════════════════════
doc.add_heading('8. 游戏状态机', level=1)

doc.add_paragraph(
    '游戏引擎维护一个有限状态机，通过 tick() 方法驱动状态转移。'
    'tick() 被 WebSocket 的 tick_loop 协程每 100ms 调用一次。'
)

doc.add_heading('状态转移图', level=2)
doc.add_paragraph(
    '                    ┌──────────────────────────────────┐\n'
    '                    │                                  │\n'
    '                    ▼                                  │\n'
    '  ┌─────────┐  start_round()  ┌───────────┐  tick   ┌───────┐\n'
    '  │ WAITING │ ───────────────→ │ COUNTDOWN │ ──────→ │ SHOOT │\n'
    '  └────▲────┘                  └───────────┘  3秒后  └───┬───┘\n'
    '       │                                                │\n'
    '       │ tick                  ┌──────────┐  下帧检测    │\n'
    '       │ (2秒后, 非自动模式)    │  RESULT  │ ←───────────┘\n'
    '       └───────────────────────┤          │\n'
    '                               └────┬─────┘\n'
    '                                    │ tick (2秒后, 自动模式)\n'
    '                                    └──→ 回到 COUNTDOWN'
)

doc.add_heading('状态说明', level=2)
add_table(doc,
    ['状态', '枚举值', '触发条件', '行为'],
    [
        ['WAITING', 'waiting', '初始化/回合结束', '等待玩家按 Start Round 或 Space 键'],
        ['COUNTDOWN', 'countdown', 'start_round() 调用', '3 秒倒计时, tick() 每秒推送剩余秒数'],
        ['SHOOT', 'shoot', '倒计时归零 (tick)', '等待下一帧检测结果, 收到后立即裁决'],
        ['RESULT', 'result', 'resolve_round() 完成', '显示结果 2 秒 (由 RESULT_DISPLAY_SECONDS 控制)'],
    ]
)

doc.add_heading('关键时间参数', level=2)
add_table(doc,
    ['参数', '值', '说明'],
    [
        ['COUNTDOWN_SECONDS', '3 秒', '倒计时持续时长'],
        ['RESULT_DISPLAY_SECONDS', '2 秒', '结果展示持续时长'],
        ['ROUND_COOLDOWN_SECONDS', '1 秒', '回合冷却 (预留)'],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════
# 9. WebSocket 协议
# ════════════════════════════════════════════
doc.add_heading('9. WebSocket 协议', level=1)

doc.add_paragraph(
    '前后端通过单一 WebSocket 端点 (ws://host:8000/ws) 进行所有实时通信。'
    '所有消息均为 JSON 格式，通过顶层 "type" 字段区分消息类型。'
)

doc.add_heading('客户端 → 服务端', level=2)
add_table(doc,
    ['type', '附加字段', '说明'],
    [
        ['frame', '{data: "<base64 JPEG>"}', '摄像头帧, ~15 FPS'],
        ['start_round', '{}', '开始新回合'],
        ['reset', '{}', '重置所有比分和历史'],
        ['get_state', '{}', '请求当前游戏状态'],
        ['toggle_auto', '{}', '切换自动对战模式'],
        ['get_stats', '{}', '请求完整统计数据'],
    ]
)

doc.add_heading('服务端 → 客户端', level=2)
add_table(doc,
    ['type', '关键字段', '触发时机'],
    [
        ['annotated_frame', 'data (base64 JPEG), predictions[]', '每收到一帧'],
        ['game_state', 'state, count/player_move/computer_move/result/reasoning', 'tick 和回合裁决'],
        ['score_update', 'player, computer, draws, round', '每回合结束后'],
        ['auto_play', 'enabled (bool)', 'toggle_auto 响应'],
        ['stats', 'wins, losses, draws, class_stats, strategy_stats, history', 'get_stats 响应'],
        ['error', 'message', '模型未加载等错误'],
        ['reset_ok', '{}', '重置确认'],
    ]
)

doc.add_heading('game_state 详情', level=2)

doc.add_paragraph('countdown 阶段:')
add_code_block(doc, """{
  "type": "game_state", "state": "countdown",
  "count": 3,    // 剩余秒数 (3, 2, 1)
  "round": 5     // 当前回合数
}""")

doc.add_paragraph('shoot 阶段:')
add_code_block(doc, """{
  "type": "game_state", "state": "shoot",
  "round": 5
}""")

doc.add_paragraph('result 阶段 (含 AI 推理详情):')
add_code_block(doc, """{
  "type": "game_state",
  "state": "result",
  "player_move": "Rock",         // 玩家检测到的手势
  "computer_move": "Paper",      // 电脑出的拳
  "predicted_player_move": "Rock", // AI 预测的玩家出拳
  "result": "lose",              // win / lose / draw
  "confidence": 0.87,            // YOLO 检测置信度
  "strategy": "markov",          // 本回合使用的策略
  "reasoning": {                 // AI 决策过程 (详见下方)
    "candidates": {
      "markov": "Rock",          // 马尔可夫预测
      "frequency": "Scissors",   // 频率分析预测
      "anti_rotate": "N/A (no rotation detected)"
    },
    "history": ["Rock","Paper","Scissors","Rock","Paper"],
    "strategy_selected": "markov",
    "predicted_player_move": "Rock",
    "computer_move": "Paper",
    "strategy_accuracies": {
      "markov": 0.667,           // 66.7% 准确率
      "frequency": 0.5,          // 50.0%
      "anti_rotate": 0.0
    }
  }
}""")

doc.add_page_break()

# ════════════════════════════════════════════
# 10. REST API 参考
# ════════════════════════════════════════════
doc.add_heading('10. REST API 参考', level=1)

doc.add_paragraph('所有接口以 /api 为前缀，返回 JSON。')

add_table(doc,
    ['方法', '路径', '说明', '响应示例'],
    [
        ['GET', '/api/health', '健康检查', '{"status":"ok", "model_loaded":true, "game_state":"waiting"}'],
        ['GET', '/api/stats', '完整统计', '{"total_games":42, "wins":18, "losses":15, "draws":9, "win_rate":0.5455, "class_stats":{...}, "strategy_stats":{...}}'],
        ['GET', '/api/history?limit=20', '最近 N 轮记录', '[{...}, ...]'],
        ['GET', '/api/state', '当前游戏状态', '{"state":"waiting", "round":5, "score":{...}}'],
        ['POST', '/api/reset', '重置所有数据', '{"status":"ok"}'],
    ]
)

doc.add_heading('/api/stats 响应字段详解', level=2)
add_table(doc,
    ['字段', '类型', '说明'],
    [
        ['total_games', 'int', '总局数 (含平局)'],
        ['wins', 'int', '玩家获胜次数'],
        ['losses', 'int', '玩家失败次数'],
        ['draws', 'int', '平局次数'],
        ['win_rate', 'float', '胜率 = wins / (wins + losses)，排除平局'],
        ['class_stats', 'dict', '{"Rock": N, "Paper": N, "Scissors": N} 玩家各类手势出拳次数'],
        ['strategy_stats', 'dict', '{"markov":{total,accuracy}, "frequency":{...}, "anti_rotate":{...}}'],
        ['score', 'dict', '{"player":N, "computer":N, "draws":N} 当前会话比分'],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════
# 11. 前端界面
# ════════════════════════════════════════════
doc.add_heading('11. 前端界面', level=1)

doc.add_paragraph(
    '前端为纯原生 HTML/CSS/JS 实现的单页应用 (SPA)，无任何框架依赖。'
    '采用暗色玻璃态主题，三栏响应式网格布局。'
)

doc.add_heading('布局结构', level=2)
doc.add_paragraph(
    '┌──────────────┬───────────────────────┬──────────────┐\n'
    '│   左栏 640px │      中栏 弹性宽度      │  右栏 340px  │\n'
    '│  Camera Feed │       Game Area       │  Statistics  │\n'
    '│              │                       │              │\n'
    '│  [摄像头画面] │  ┌── Scoreboard ──┐   │  Win Rate    │\n'
    '│              │  │ You 5 vs 3 AI  │   │  [Gauge 环]  │\n'
    '│              │  └────────────────┘   │              │\n'
    '│              │  ┌── Game Display ─┐  │  W/L/D Bars  │\n'
    '│              │  │  3...2...1 倒计时│  │              │\n'
    '│              │  │  ✊ VS ✋      │  │  Move Dist   │\n'
    '│              │  │  Computer Wins!│  │              │\n'
    '│              │  └────────────────┘   │  Strategy %  │\n'
    '│              │                       │              │\n'
    '│              │  [Start Round] [Auto] │  History      │\n'
    '│              │                       │  Table        │\n'
    '│              │  ┌── AI Console ──┐   │              │\n'
    '│              │  │ Candidates...  │   │  [Reset All] │\n'
    '│              │  │ Accuracy...    │   │              │\n'
    '│              │  │ Decision...    │   │              │\n'
    '│              │  └────────────────┘   │              │\n'
    '└──────────────┴───────────────────────┴──────────────┘'
)

doc.add_heading('前端 JavaScript 模块', level=2)
add_table(doc,
    ['模块', '文件', '类/职责'],
    [
        ['WebSocket 客户端', 'websocket.js', 'WebSocketClient 类: 连接管理、自动重连(指数退避)、消息分发'],
        ['摄像头管理', 'main.js', 'CameraManager 类: getUserMedia、帧捕获(15FPS)、base64 编码'],
        ['应用入口', 'main.js', 'App 类: 初始化 WS/摄像头/游戏/统计、绑定 annotated_frame 处理'],
        ['游戏 UI', 'game.js', 'GameUI 类: 状态渲染(倒计时/结果/AI控制台)、按钮事件、键盘快捷键'],
        ['统计面板', 'stats.js', 'StatsUI 类: Canvas 环形仪表盘、W/L/D 进度条、策略准确率、历史表格'],
    ]
)

doc.add_heading('键盘快捷键', level=2)
add_table(doc,
    ['按键', '功能'],
    [
        ['Space', '开始新回合 (等同于点击 Start Round)'],
        ['A', '切换自动对战模式 (等同于点击 Auto)'],
    ]
)

doc.add_heading('配色方案', level=2)
add_table(doc,
    ['用途', '色值', '说明'],
    [
        ['主背景', '#0A0A14', '深暗蓝黑'],
        ['面板底色', 'rgba(22,22,58,0.85)', '半透明玻璃态'],
        ['主题色', '#E94560', '霓虹红'],
        ['石头色', '#4FC3F7', '浅蓝'],
        ['布色', '#81C784', '绿色'],
        ['剪刀色', '#EF5350', '红色'],
        ['胜', '#4CAF50', '绿色'],
        ['负', '#F44336', '红色'],
        ['平', '#FF9800', '橙色'],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════
# 12. 数据库设计
# ════════════════════════════════════════════
doc.add_heading('12. 数据库设计', level=1)

doc.add_paragraph(
    '使用 SQLite 作为本地持久化存储，数据库文件位于 backend/game_history.db，'
    '在 backend/database.py 首次导入时自动创建表和索引。'
    '使用 WAL 日志模式以支持并发读写。'
)

doc.add_heading('rounds 表结构', level=2)
add_table(doc,
    ['字段', '类型', '约束', '说明'],
    [
        ['id', 'INTEGER', 'PRIMARY KEY AUTOINCREMENT', '自增主键'],
        ['round_number', 'INTEGER', 'NOT NULL', '回合编号'],
        ['player_move', 'TEXT', 'NOT NULL', '玩家检测手势 (Rock/Paper/Scissors)'],
        ['player_confidence', 'REAL', 'NULLABLE', 'YOLO 检测置信度 (0-1)'],
        ['computer_move', 'TEXT', 'NOT NULL', '电脑出拳 (Rock/Paper/Scissors)'],
        ['predicted_player_move', 'TEXT', 'NULLABLE', 'AI 预测的玩家出拳'],
        ['result', 'TEXT', 'NOT NULL', '玩家结果 (win/lose/draw)'],
        ['strategy_used', 'TEXT', 'NULLABLE', '本回合使用的策略名 (markov/frequency/anti_rotate/fallback)'],
        ['timestamp', 'DATETIME', 'DEFAULT CURRENT_TIMESTAMP', '记录时间戳'],
    ]
)

doc.add_heading('索引', level=2)
add_table(doc,
    ['索引名', '字段', '用途'],
    [
        ['idx_rounds_result', 'result', '加速按结果聚合查询'],
        ['idx_rounds_timestamp', 'timestamp', '加速按时间排序查询'],
    ]
)

doc.add_heading('关键查询', level=2)

doc.add_paragraph('总体统计 (get_stats):')
add_code_block(doc, """SELECT COUNT(*) FROM rounds                               -- total_games
SELECT COUNT(*) FROM rounds WHERE result = 'win'          -- wins
SELECT COUNT(*) FROM rounds WHERE result = 'lose'         -- losses
SELECT COUNT(*) FROM rounds WHERE result = 'draw'         -- draws""")

doc.add_paragraph('玩家出拳分布 (get_class_stats):')
add_code_block(doc, """SELECT player_move, COUNT(*) as cnt
FROM rounds GROUP BY player_move""")

doc.add_paragraph('最近 N 局 (get_recent_rounds):')
add_code_block(doc, """SELECT * FROM rounds ORDER BY id DESC LIMIT ?""")

doc.add_page_break()

# ════════════════════════════════════════════
# 13. 配置文件
# ════════════════════════════════════════════
doc.add_heading('13. 配置文件', level=1)

doc.add_heading('backend/config.py — 核心配置', level=2)
add_table(doc,
    ['常量', '默认值', '类别', '说明'],
    [
        ['MODEL_PATH', 'models/rps_yolov8n.pt', '路径', 'YOLO 模型文件路径'],
        ['DATABASE_PATH', 'backend/game_history.db', '路径', 'SQLite 数据库路径'],
        ['COUNTDOWN_SECONDS', '3', '游戏', '倒计时秒数'],
        ['RESULT_DISPLAY_SECONDS', '2', '游戏', '结果显示秒数'],
        ['ROUND_COOLDOWN_SECONDS', '1', '游戏', '回合冷却秒数'],
        ['CONFIDENCE_THRESHOLD', '0.3', '检测', 'YOLO 最小置信度阈值'],
        ['CAMERA_WIDTH', '640', '摄像头', '摄像头宽度'],
        ['CAMERA_HEIGHT', '640', '摄像头', '摄像头高度'],
        ['CAMERA_FPS', '30', '摄像头', '摄像头帧率'],
        ['WS_FRAME_QUALITY', '85', '通信', 'JPEG 压缩质量 (0-100)'],
        ['MARKOV_ORDER', '3', 'AI', '马尔可夫链阶数'],
        ['HISTORY_WINDOW', '50', 'AI', '历史窗口 (预留)'],
        ['FREQUENCY_DECAY', '0.85', 'AI', '频率衰减因子'],
        ['RECENT_WINDOW', '10', 'AI', '推理显示历史长度'],
        ['META_WINDOW', '20', 'AI', '元策略滑动窗口'],
    ]
)

doc.add_heading('data.yaml — 数据集配置', level=2)
add_code_block(doc, """# YOLO dataset config - Rock Paper Scissors
path: datasets/          # 数据集根目录
train: train/images      # 训练集图像路径 (相对 path)
val: val/images          # 验证集图像路径
test: test/images        # 测试集图像路径

nc: 3                    # 类别数量
names:                   # 类别名称映射
  0: Rock
  1: Paper
  2: Scissors""")

doc.add_page_break()

# ════════════════════════════════════════════
# 14. 部署与运行
# ════════════════════════════════════════════
doc.add_heading('14. 部署与运行', level=1)

doc.add_heading('环境要求', level=2)
add_table(doc,
    ['项目', '要求'],
    [
        ['操作系统', 'Windows 10+ / Linux / macOS'],
        ['Python', '3.10 或更高'],
        ['CUDA (可选)', 'CUDA 11.8+ (GPU 训练/推理)'],
        ['浏览器', '支持 WebSocket 和 getUserMedia 的现代浏览器 (Chrome/Firefox/Edge)'],
        ['摄像头', 'USB/内置摄像头, 建议 640×480 以上分辨率'],
    ]
)

doc.add_heading('完整部署流程', level=2)
doc.add_paragraph('步骤 1: 克隆项目并安装依赖')
add_code_block(doc, """cd Rock-Paper-Scissors
pip install -r requirements.txt""")

doc.add_paragraph('步骤 2: 准备数据集 (如已有标注数据)')
add_code_block(doc, """python scripts/prepare_data.py""")

doc.add_paragraph('步骤 3: 训练 YOLO 模型')
add_code_block(doc, """# GPU 训练 (推荐)
python scripts/train.py --epochs 100 --batch 64 --device cuda

# CPU 训练 (备选)
python scripts/train.py --epochs 10 --batch 4 --device cpu""")

doc.add_paragraph('步骤 4: 启动游戏服务')
add_code_block(doc, """python backend/app.py""")

doc.add_paragraph('步骤 5: 打开浏览器')
add_code_block(doc, """http://localhost:8000""")

doc.add_heading('Windows 注意事项', level=2)
doc.add_paragraph(
    '• DataLoader workers 必须设为 0 (避免 multiprocessing 问题)\n'
    '• AMP 自动混合精度在 Windows 上需关闭 (amp=False)\n'
    '• 首次启动浏览器会请求摄像头权限，需允许'
)

doc.add_page_break()

# ════════════════════════════════════════════
# 15. 依赖清单
# ════════════════════════════════════════════
doc.add_heading('15. 依赖清单', level=1)

doc.add_paragraph('完整依赖列表 (requirements.txt):')
add_table(doc,
    ['包名', '版本要求', '用途'],
    [
        ['ultralytics', '≥8.0.0', 'YOLOv8 训练和推理框架'],
        ['torch', '≥2.0.0', 'PyTorch 深度学习框架'],
        ['opencv-python', '≥4.8.0', '图像处理和编解码'],
        ['fastapi', '≥0.104.0', 'Web API 框架'],
        ['uvicorn[standard]', '≥0.24.0', 'ASGI 服务器'],
        ['websockets', '≥12.0', 'WebSocket 协议支持'],
        ['scikit-learn', '≥1.3.0', '数据集分层划分'],
        ['numpy', '≥1.24.0', '数值计算'],
        ['Pillow', '≥10.0.0', '图像处理 (YOLO 依赖)'],
        ['python-multipart', '≥0.0.6', 'FastAPI 表单数据解析'],
    ]
)

doc.add_paragraph(
    f'\n文档版本: 1.0\n'
    f'生成日期: {datetime.date.today().isoformat()}\n'
    f'项目路径: {BASE_DIR}'
)

# ── Save ──
doc.save(str(OUTPUT_PATH))
print(f"Done: {OUTPUT_PATH}")
print(f"Size: {OUTPUT_PATH.stat().st_size / 1024:.1f} KB")
